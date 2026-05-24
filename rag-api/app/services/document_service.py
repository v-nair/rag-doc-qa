import csv
import io
import logging
import os
import uuid
import zipfile

import fitz  # pymupdf
from docx import Document
from openpyxl import load_workbook

from config import CHUNK_SIZE, CHUNK_OVERLAP, MAX_IMAGES_PER_DOC, SUPPORTED_EXTENSIONS
from services.vision_service import describe_image, is_image_meaningful

logger = logging.getLogger(__name__)


class UnsupportedFileType(ValueError):
    pass


def parse_and_chunk(file_bytes: bytes, filename: str) -> tuple[str, list[str]]:
    """Parse a file by extension and return (doc_id, list_of_text_chunks).

    Supported: pdf, docx, xlsx, csv, png, jpg/jpeg.
    PDF and DOCX images are described via GPT-4o vision and inlined as text.
    Standalone images become a single description chunk.
    """
    doc_id = str(uuid.uuid4())
    ext = os.path.splitext(filename)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(
            f"Unsupported file type: {ext}. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext == ".pdf":
        text = _extract_pdf(file_bytes)
    elif ext == ".docx":
        text = _extract_docx(file_bytes)
    elif ext == ".xlsx":
        text = _extract_xlsx(file_bytes)
    elif ext == ".csv":
        text = _extract_csv(file_bytes)
    else:  # .png, .jpg, .jpeg
        text = _extract_image(file_bytes, ext)

    chunks = _split_text(text)
    return doc_id, chunks


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from every page plus vision descriptions of embedded images."""
    sections: list[str] = []
    image_count = 0

    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text().strip()
            if page_text:
                sections.append(f"[Page {page_num}]\n{page_text}")

            for img_info in page.get_images(full=True):
                if image_count >= MAX_IMAGES_PER_DOC:
                    break
                xref = img_info[0]
                try:
                    base = doc.extract_image(xref)
                    img_bytes = base["image"]
                    ext = base.get("ext", "png")
                except Exception as e:
                    logger.warning(f"Failed to extract PDF image xref={xref}: {e}")
                    continue

                if not is_image_meaningful(img_bytes):
                    continue

                description = describe_image(img_bytes, mime=f"image/{ext}")
                if description:
                    sections.append(f"[Image on page {page_num}]\n{description}")
                    image_count += 1

            if image_count >= MAX_IMAGES_PER_DOC:
                logger.info(f"Reached image cap ({MAX_IMAGES_PER_DOC}), skipping the rest")
                break

    return "\n\n".join(sections)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract paragraph + table text from DOCX, plus vision descriptions of media images."""
    sections: list[str] = []

    document = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    if paragraphs:
        sections.append("\n".join(paragraphs))

    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        if rows:
            sections.append("[Table]\n" + "\n".join(rows))

    image_count = 0
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
            for name in media:
                if image_count >= MAX_IMAGES_PER_DOC:
                    break
                img_bytes = z.read(name)
                if not is_image_meaningful(img_bytes):
                    continue
                ext = os.path.splitext(name)[1].lstrip(".") or "png"
                description = describe_image(img_bytes, mime=f"image/{ext}")
                if description:
                    sections.append(f"[Image: {os.path.basename(name)}]\n{description}")
                    image_count += 1
    except zipfile.BadZipFile:
        logger.warning("DOCX is not a valid zip; skipping image extraction")

    return "\n\n".join(sections)


def _extract_xlsx(file_bytes: bytes) -> str:
    """Flatten every sheet's rows into pipe-separated text."""
    sections: list[str] = []
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows: list[str] = []
            for row in ws.iter_rows(values_only=True):
                if any(c is not None for c in row):
                    rows.append(" | ".join("" if c is None else str(c) for c in row))
            if rows:
                sections.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
    finally:
        wb.close()
    return "\n\n".join(sections)


def _extract_csv(file_bytes: bytes) -> str:
    text = _decode_text(file_bytes)
    reader = csv.reader(io.StringIO(text))
    rows = [" | ".join(row) for row in reader if any(cell.strip() for cell in row)]
    return "\n".join(rows)


def _extract_image(file_bytes: bytes, ext: str) -> str:
    """Describe a single standalone image via GPT-4o vision."""
    mime_ext = "jpeg" if ext in (".jpg", ".jpeg") else ext.lstrip(".")
    description = describe_image(file_bytes, mime=f"image/{mime_ext}")
    if not description:
        raise ValueError("Could not describe image (vision API failed)")
    return f"[Image]\n{description}"


def _decode_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _split_text(text: str) -> list[str]:
    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = start + CHUNK_SIZE
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks
